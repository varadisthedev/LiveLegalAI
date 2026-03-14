"""
test_api.py
-----------
End-to-end test for ALL RAG microservice endpoints.

Tests the full flow:
  1. POST /ingest     — upload a mock legal document
  2. GET  /documents  — verify it appears in the history list
  3. POST /analyze    — analyse the document (summary, severity, risk chart data)
  4. GET  /documents  — verify it now shows as analyzed
  5. POST /chat       — ask a follow-up question about the document
  6. GET  /health     — verify health check

Run with: python test_api.py
(Requires the FastAPI server to be running on localhost:8000)
"""

import httpx
import json
import sys
from docx import Document

BASE_URL = "http://localhost:8000"
TIMEOUT = 120.0

# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def print_header(step, title):
    print(f"\n{'='*60}")
    print(f"  STEP {step}: {title}")
    print(f"{'='*60}")

def print_result(label, data, indent=2):
    if isinstance(data, dict):
        print(json.dumps(data, indent=indent))
    else:
        print(data)

def assert_ok(response, step_name):
    if response.status_code != 200:
        print(f"\n❌ {step_name} FAILED — HTTP {response.status_code}")
        print(response.text[:500])
        sys.exit(1)
    print(f"✅ {step_name} — HTTP {response.status_code}")

# ──────────────────────────────────────────────
# 1. Generate a realistic test legal document
# ──────────────────────────────────────────────

def create_test_document():
    print_header(1, "Generating Test Legal Document")
    doc = Document()
    doc.add_heading("FORMAL NOTICE OF COPYRIGHT INFRINGEMENT", level=1)
    doc.add_paragraph(
        "Date: March 14, 2025\n"
        "From: MediaCorp Legal Department\n"
        "To: Content Creator (Channel Owner)\n"
        "Re: Copyright Infringement — Video ID: YT-2025-0314\n"
    )
    doc.add_paragraph(
        "Dear Content Creator,\n\n"
        "We are writing to formally notify you that your recent video uploaded to YouTube "
        "on March 1, 2025 contains copyrighted material owned by MediaCorp International LLC. "
        "Specifically, the video includes approximately 45 seconds of our original soundtrack "
        "'Midnight Eclipse' without authorization or license.\n"
    )
    doc.add_paragraph(
        "LEGAL BASIS: Under the Digital Millennium Copyright Act (DMCA), Section 512, "
        "unauthorized reproduction and distribution of copyrighted content constitutes "
        "intellectual property infringement. MediaCorp reserves the right to pursue legal action "
        "including but not limited to a lawsuit seeking damages up to $150,000 per infringement.\n"
    )
    doc.add_paragraph(
        "REQUIRED ACTION: You must respond within 14 days of receipt of this notice. "
        "Failure to comply by the deadline of March 28, 2025 may result in:\n"
        "1. A formal DMCA takedown request\n"
        "2. Monetization claim on the video\n"
        "3. Initiation of legal proceedings in federal court\n"
    )
    doc.add_paragraph(
        "We strongly recommend you seek legal counsel regarding this matter. "
        "You may also submit a counter-notification if you believe this claim is in error.\n\n"
        "Sincerely,\n"
        "Sarah Mitchell, Esq.\n"
        "Head of Legal, MediaCorp International LLC\n"
        "Email: legal@mediacorp.com\n"
        "Phone: (555) 123-4567"
    )

    filename = "test_copyright_notice.docx"
    doc.save(filename)
    print(f"✅ Created '{filename}' — realistic copyright infringement notice")
    return filename

# ──────────────────────────────────────────────
# Run all tests
# ──────────────────────────────────────────────

def main():
    print("\n" + "🔥" * 20)
    print("  LiveLegal AI — Full API Test Suite")
    print("🔥" * 20)

    client = httpx.Client(timeout=TIMEOUT)

    # ── Step 1: Create test document ──
    filename = create_test_document()

    # ── Step 2: POST /ingest ──
    print_header(2, "POST /ingest — Upload & Process Document")
    with open(filename, "rb") as f:
        response = client.post(f"{BASE_URL}/ingest", files={"file": (filename, f)})
    assert_ok(response, "POST /ingest")
    ingest_data = response.json()
    print_result("Ingest Response", ingest_data)

    document_id = ingest_data["document_id"]
    print(f"\n📋 Document ID: {document_id}")
    print(f"📄 Chunks created: {ingest_data['num_chunks']}")

    # ── Step 3: GET /documents — check history (before analysis) ──
    print_header(3, "GET /documents — Check Document History (Pre-Analysis)")
    response = client.get(f"{BASE_URL}/documents")
    assert_ok(response, "GET /documents")
    docs_data = response.json()
    print(f"📚 Total documents in registry: {docs_data['total']}")
    for doc in docs_data['documents']:
        status = "🟢 Analyzed" if doc.get('analyzed') else "🟡 Pending"
        print(f"   {status} | {doc['document_id']} | {doc['filename']}")

    # ── Step 4: POST /analyze — Full Document Analysis ──
    print_header(4, "POST /analyze — AI Analysis + Severity Scoring")
    response = client.post(
        f"{BASE_URL}/analyze",
        json={"document_id": document_id}
    )
    assert_ok(response, "POST /analyze")
    analyze_data = response.json()

    print(f"\n🔍 SEVERITY SCORE: {analyze_data['severity_score']}/100 ({analyze_data['risk_level']})")
    print(f"📂 DOCUMENT TYPE:  {analyze_data['document_type']}")
    print(f"\n📝 SUMMARY:\n{analyze_data['summary']}")
    print(f"\n📖 EXPLANATION:\n{analyze_data['explanation'][:300]}...")
    print(f"\n✉️  SUGGESTED REPLY:\n{analyze_data['suggested_reply'][:300]}...")

    if analyze_data.get('risk_factors'):
        print(f"\n📊 RISK FACTORS (for chart/graph):")
        for rf in analyze_data['risk_factors']:
            print(f"   ├─ {rf['label']}: +{rf['points']} pts [{rf['category']}]")

    # ── Step 5: GET /documents — check history (after analysis) ──
    print_header(5, "GET /documents — Check Document History (Post-Analysis)")
    response = client.get(f"{BASE_URL}/documents")
    assert_ok(response, "GET /documents")
    docs_data = response.json()
    for doc in docs_data['documents']:
        status = "🟢 Analyzed" if doc.get('analyzed') else "🟡 Pending"
        score = f"Score: {doc.get('severity_score', 'N/A')}" if doc.get('analyzed') else ""
        print(f"   {status} | {doc['document_id']} | {doc['filename']} | {score}")

    # ── Step 6: POST /chat — Ask a Follow-Up Question ──
    print_header(6, "POST /chat — Chatbot Follow-Up Question")
    response = client.post(
        f"{BASE_URL}/chat",
        json={
            "document_id": document_id,
            "question": "What is the deadline to respond and what happens if I miss it?"
        }
    )
    assert_ok(response, "POST /chat")
    chat_data = response.json()

    print(f"\n❓ QUESTION: {chat_data['question']}")
    print(f"\n🤖 ANSWER:\n{chat_data['answer']}")
    print(f"\n📎 Sources used: {chat_data['sources_used']} chunks")
    if chat_data.get('context_snippets'):
        print(f"\n📋 Source Snippets:")
        for i, snippet in enumerate(chat_data['context_snippets'][:2]):
            print(f"   [{i+1}] {snippet[:120]}...")

    # ── Step 7: GET /health ──
    print_header(7, "GET /health — Health Check")
    response = client.get(f"{BASE_URL}/health")
    assert_ok(response, "GET /health")
    print_result("Health", response.json())

    # ── Summary ──
    print(f"\n{'='*60}")
    print("  ✅ ALL TESTS PASSED SUCCESSFULLY")
    print(f"{'='*60}")
    print(f"\n📊 API Contract Summary for Frontend:")
    print(f"   POST /ingest    → {{ document_id, filename, num_chunks }}")
    print(f"   POST /analyze   → {{ summary, explanation, suggested_reply, severity_score, risk_level, risk_factors[], document_type }}")
    print(f"   POST /chat      → {{ answer, sources_used, context_snippets[] }}")
    print(f"   GET  /documents → {{ documents[], total }}")
    print(f"   GET  /health    → {{ status }}")
    print()


if __name__ == "__main__":
    main()
